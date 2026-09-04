"""把支持的 document 文件统一提取成纯文本。

这一层只负责“读文件内容”，不负责 embedding、分类或索引。
不同格式的文件最后都会返回同一种 ExtractionResult，后面的分段和
向量化代码不需要知道原文件是 PDF 还是 DOCX。
"""

import concurrent.futures
import re
from dataclasses import asdict, dataclass
from pathlib import Path
from typing import Optional


# 单个文件提取的超时时间。个别下载损坏的文件（比如把 HTTP 错误响应存成了 .pdf）
# 会让 pypdf 的容错解析陷入极慢的恢复尝试，不设超时的话可能拖住整批索引。
EXTRACTION_TIMEOUT_SECONDS = 30


@dataclass
class ExtractionResult:
    """一次文本提取的统一结果。

    status 的可能值：
    - success: 成功提取到文本
    - no_text: 文件可读取，但没有提取到文字，例如扫描版 PDF
    - unsupported: 当前版本没有对应的提取器
    - error: 提取过程中发生异常
    """

    path: str
    status: str
    text: str = ""
    extractor: Optional[str] = None
    error: Optional[str] = None

    @property
    def char_count(self):
        """返回清洗后文本的字符数，方便报告统计。"""

        return len(self.text)

    def to_dict(self):
        """转成普通字典，后续可以直接写入 JSON 或 SQLite。"""

        result = asdict(self)
        result["char_count"] = self.char_count
        return result


# 这里列的是“当前提取器支持的后缀”，不等同于 Week 1 的全部 document 类型。
# .xls（老版二进制格式）和 .xlsx（新版 XML 容器格式）不是同一个格式，不能用同一个
# 库：openpyxl 只读 .xlsx，读不了老版二进制格式；能读老版 .xls 的库是 xlrd，但
# xlrd 2.0 之后反而砍掉了 .xlsx 支持（社区当年在解析 xlsx 那块出过安全漏洞），
# 两个库分工互补，不能互相替代，所以都装了。.doc（老版二进制 Word 格式，同样不是
# .docx）也是类似情况：python-docx 只认新版 XML 格式，读不了老版二进制格式，
# 见 _extract_doc 的注释。
SUPPORTED_EXTENSIONS = {
    ".pdf",
    ".docx",
    ".doc",
    ".txt",
    ".md",
    ".markdown",
    ".html",
    ".htm",
    ".csv",
    ".json",
    ".xlsx",
    ".xls",
    ".log",
    # 零散的源码/脚本/配置文件（不在任何 git 项目目录里的，项目内部文件走
    # Stage-project 整体处理，不会走到这里）——本身就是纯文本，跟 TXT/MD 用
    # 同一个提取器即可，不用单独写代码。覆盖范围不只看当前用户自己的技术栈
    # （前端 + Python），常见后端/移动端语言也一并覆盖，避免只服务一种开发者。
    ".py", ".js", ".jsx", ".ts", ".tsx", ".vue", ".svelte",
    ".css", ".scss", ".sass", ".less",
    ".java", ".go", ".rs", ".rb", ".php", ".c", ".h", ".cpp", ".hpp",
    ".cs", ".swift", ".kt", ".m", ".mm",
    ".sh", ".bash", ".zsh", ".ps1", ".bat",
    ".yaml", ".yml", ".toml", ".ini", ".cfg", ".sql", ".xml",
}


def _dispatch_extract(suffix, path):
    """按后缀分发到对应的提取函数，返回 (text, extractor_name)。"""

    if suffix == ".pdf":
        return _extract_pdf(path), "pypdf"
    elif suffix == ".docx":
        return _extract_docx(path), "python-docx"
    elif suffix == ".doc":
        return _extract_doc(path), "olefile"
    elif suffix in {".html", ".htm"}:
        return _extract_html(path), "beautifulsoup4"
    elif suffix == ".csv":
        return _extract_csv(path), "python-csv"
    elif suffix == ".json":
        return _extract_json(path), "python-json"
    elif suffix == ".xlsx":
        return _extract_xlsx(path), "openpyxl"
    elif suffix == ".xls":
        return _extract_xls(path)
    else:
        # 源码/脚本/配置/日志这些零散文本文件，全部落到这里，跟 TXT/MD 用同一个
        # 提取器——它们本质上都只是"纯文本"，不需要专门的解析逻辑。
        return _extract_plain_text(path), "python-text"


def extract_text(file_path):
    """提取一个文件的纯文本，并返回统一的 ExtractionResult。

    单个文件失败不会在这里向外抛出异常，而是返回 status="error" 和
    error 字段。这样批量处理时可以记录失败文件，继续处理其他文件。
    """

    path = Path(file_path).expanduser()
    suffix = path.suffix.lower()

    if suffix not in SUPPORTED_EXTENSIONS:
        return ExtractionResult(
            path=str(path),
            status="unsupported",
            error=f"暂不支持的文件格式: {suffix or '(无扩展名)'}",
        )

    # 提取放到独立线程里跑，主线程用 future.result(timeout=...) 掐表等。
    # 注意：这里故意每次新建一个只有 1 个 worker 的 executor，而不是复用一个共享池——
    # 万一某份文件真的卡死不返回，我们只是不再等它（shutdown(wait=False)，不阻塞退出），
    # 那个卡住的线程会变成孤儿线程留在后台，但不会占用下一份文件要用的 worker 名额。
    # 如果复用一个固定大小的共享池，一旦某个 worker 被卡死的任务占住，就永久少了一条
    # 处理后续文件的产能，只有重启进程才能恢复。
    executor = concurrent.futures.ThreadPoolExecutor(max_workers=1)
    future = executor.submit(_dispatch_extract, suffix, path)
    try:
        text, extractor = future.result(timeout=EXTRACTION_TIMEOUT_SECONDS)
        executor.shutdown(wait=False)
    except concurrent.futures.TimeoutError:
        executor.shutdown(wait=False)
        return ExtractionResult(
            path=str(path),
            status="error",
            extractor=None,
            error=f"提取超时（超过 {EXTRACTION_TIMEOUT_SECONDS} 秒），文件可能已损坏",
        )
    except Exception as exc:  # 单文件异常转为记录，不阻塞整批索引
        executor.shutdown(wait=False)
        return ExtractionResult(
            path=str(path),
            status="error",
            extractor=None,
            error=f"{type(exc).__name__}: {exc}",
        )

    text = _normalize_text(text)
    status = "success" if text else "no_text"
    return ExtractionResult(
        path=str(path),
        status=status,
        text=text,
        extractor=extractor,
    )


def _decode_bytes(raw):
    """按常见编码顺序尝试解码：UTF-8（含 BOM）优先，中文语料常见的 GB18030 兜底。
    TXT/MD/CSV 三种纯文本格式共用同一套解码逻辑，不重复写三遍。
    """

    for encoding in ("utf-8-sig", "utf-8", "gb18030"):
        try:
            return raw.decode(encoding)
        except UnicodeDecodeError:
            continue

    raise UnicodeDecodeError(
        "text",
        raw,
        0,
        len(raw),
        "无法按 UTF-8 或 GB18030 解码",
    )


def _extract_plain_text(path):
    """读取 TXT/MD，并兼容常见的 UTF-8、GB18030 编码。"""

    return _decode_bytes(path.read_bytes())


def _extract_csv(path):
    """把 CSV 每一行提取出来，单元格之间用制表符连接——跟 _extract_docx 里处理
    表格的约定一致，保留列与列之间的边界，后续检索依然能命中单元格里的关键词，
    只是不做更复杂的表格结构化（比如按表头对齐列名），留到后续阶段。
    """

    import csv
    import io

    text = _decode_bytes(path.read_bytes())
    rows = []
    for row in csv.reader(io.StringIO(text)):
        cells = [cell.strip() for cell in row]
        if any(cells):
            rows.append("\t".join(cells))
    return "\n".join(rows)


# xlsx/xls 提取多工作表时，用来标注"这是哪个工作表"的分隔行格式——下面三处
# 提取函数（openpyxl 正常路径、xlrd 老 OLE2 路径、SpreadsheetML XML 路径）
# 都要写这个标记，chunking.py 的 chunk_tabular_text() 要靠认出它来切表格块。
# 定义只放这一处，chunking.py 从这里 import 过去拼正则，不在两个文件里
# 各写一份一样的字符串（改一处忘了改另一处，两边就对不上了）。
SHEET_MARKER_PREFIX = "[sheet: "
SHEET_MARKER_SUFFIX = "]"


def format_sheet_marker(sheet_name):
    return f"{SHEET_MARKER_PREFIX}{sheet_name}{SHEET_MARKER_SUFFIX}"


def _extract_xlsx(path):
    """提取 XLSX 每个工作表的单元格内容，单元格之间用制表符连接（跟 CSV/DOCX
    表格提取用同一套约定）。data_only=True 读公式单元格最后一次计算出来的缓存值，
    不是公式文本本身（比如读到"1200"而不是"=SUM(A1:A10)"，这才是用户真正关心的
    内容——公式表达式本身没有语义检索价值）。read_only=True 用流式读取，避免大
    表格整个加载进内存。多个工作表的内容分开标注工作表名，不糊成一整段。
    """

    from openpyxl import load_workbook

    workbook = load_workbook(str(path), data_only=True, read_only=True)
    try:
        blocks = []
        for sheet in workbook.worksheets:
            sheet_rows = []
            for row in sheet.iter_rows(values_only=True):
                cells = [str(cell).strip() if cell is not None else "" for cell in row]
                if any(cells):
                    sheet_rows.append("\t".join(cells))
            if sheet_rows:
                blocks.append(f"{format_sheet_marker(sheet.title)}\n" + "\n".join(sheet_rows))
        return "\n\n".join(blocks)
    finally:
        workbook.close()


def _extract_xls(path):
    """.xls 这个扩展名不完全可信——真实语料（~/Downloads）里抽查发现，同一个 .xls
    后缀下真实混着三种不同格式：真正的老版二进制格式（OLE2）、被错误存成 .xls
    扩展名的 xlsx（zip 容器，Excel 打开不挑扩展名，一些导出工具就顺手存错了）、
    以及 SpreadsheetML（Excel 2003 XML 电子表格格式，历史报表/BI 系统的常见导出
    格式，同样是 Excel 能打开但既不是 xls 也不是 xlsx）。这不是预防性设计，是
    真实抽查 20 份 .xls 后发现一半以上其实是 SpreadsheetML 才加的分流逻辑——
    必须看文件开头的真实字节判断格式，不能只信扩展名。返回 (text, 实际用的
    提取器名字)，提取器名字如实反映真正走了哪条分支，不是写死的 "xlrd"。
    """

    header = path.read_bytes()[:16]
    if header.startswith(b"PK"):
        return _extract_xlsx(path), "openpyxl(伪装成xls的xlsx)"
    if header.lstrip()[:5] == b"<?xml":
        return _extract_spreadsheetml(path), "xml.etree(SpreadsheetML)"
    return _extract_xls_binary(path), "xlrd"


_SPREADSHEETML_NS = "urn:schemas-microsoft-com:office:spreadsheet"


#  合法 XML 里裸的 "&" 必须转义成 "&amp;"（因为 & 是实体引用的起始符），但真实
# 语料里抓到过好几份文件的公司名带 "&"（比如 `"D&G Kompiuteriai"`），生成这份
# "XML" 的工具当年没做转义，导致文件本身就不是合法 XML，标准解析器直接报错拒绝
# 解析。这是源数据的问题，不是我们代码的问题，但既然内容本身有真实价值（公司名），
# 值得在解析前做一次"抢救性"预处理：把不属于任何合法实体引用的裸 "&" 都转义成
# "&amp;" 再交给解析器，合法的实体引用（&amp; &lt; &gt; &quot; &apos; 以及
# &#123; 这种数字实体）不动，避免把本来就合法的转义符号"二次转义"坏成 "&amp;amp;"。
_BARE_AMPERSAND_PATTERN = re.compile(r"&(?!amp;|lt;|gt;|quot;|apos;|#\d+;|#x[0-9a-fA-F]+;)")


def _extract_spreadsheetml(path):
    """提取 SpreadsheetML（Excel 2003 XML 电子表格格式）——纯 XML 结构
    （Workbook > Worksheet > Table > Row > Cell > Data），标准库 xml.etree
    就能读，不需要再装专门的库。"""

    import xml.etree.ElementTree as ET

    ns = {"ss": _SPREADSHEETML_NS}
    raw_text = _decode_bytes(path.read_bytes())
    fixed_text = _BARE_AMPERSAND_PATTERN.sub("&amp;", raw_text)
    root = ET.fromstring(fixed_text)

    blocks = []
    for worksheet in root.findall("ss:Worksheet", ns):
        sheet_name = worksheet.get(f"{{{_SPREADSHEETML_NS}}}Name", "")
        sheet_rows = []
        for row in worksheet.findall("ss:Table/ss:Row", ns):
            cells = []
            for cell in row.findall("ss:Cell", ns):
                data = cell.find("ss:Data", ns)
                cells.append((data.text or "").strip() if data is not None else "")
            if any(cells):
                sheet_rows.append("\t".join(cells))
        if sheet_rows:
            blocks.append(f"{format_sheet_marker(sheet_name)}\n" + "\n".join(sheet_rows))
    return "\n\n".join(blocks)


def _extract_xls_binary(path):
    """提取真正的老版二进制 XLS（OLE2 格式），用 xlrd 读——它是唯一还在维护对
    这个老格式支持的库（2.0 版本之后反而砍掉了 xlsx 支持，见 SUPPORTED_EXTENSIONS
    的注释）。

    日期类型单元格要特殊处理：xlrd 读出来的原始值是 Excel 内部的浮点序列号
    （比如 45678.0），不是人能看懂的日期字符串，要用 cell.ctype 先判断这个单元格
    是不是日期类型，是的话再用 xldate_as_datetime 配合 book.datemode 转成真正
    的日期，不然存进数据库/喂给 LLM 的都是一串没有意义的数字。
    """

    import xlrd

    book = xlrd.open_workbook(str(path))
    blocks = []
    for sheet in book.sheets():
        sheet_rows = []
        for row_idx in range(sheet.nrows):
            cells = []
            for cell in sheet.row(row_idx):
                if cell.ctype == xlrd.XL_CELL_EMPTY:
                    value = ""
                elif cell.ctype == xlrd.XL_CELL_DATE:
                    try:
                        value = xlrd.xldate_as_datetime(cell.value, book.datemode)
                    except Exception:
                        value = cell.value  # 极少数畸形日期值转换失败，退回原始数字，不中断整份提取
                else:
                    value = cell.value
                cells.append(str(value).strip())
            if any(cells):
                sheet_rows.append("\t".join(cells))
        if sheet_rows:
            blocks.append(f"{format_sheet_marker(sheet.name)}\n" + "\n".join(sheet_rows))
    return "\n\n".join(blocks)


# 顶层是数组、元素数量超过这个阈值的 JSON，判定为"很多条独立记录堆在一个数组里"
# （聊天记录导出、日志导出、邮件导出都是这个形状），提取时均匀抽样，不整份塞进
# full_text——真实撞过的坑：一份 6000 万字符的聊天记录导出文件，顶层数组里塞了
# 几百上千条互相独立的对话，被当成一整篇连续文章处理会导致两个问题同时发生：
# (1) 切分之后 chunk 数量爆炸（这一份文件真实切出过 9 万+ chunk），向量化开销
# 失控；(2) 打标签阶段只读 full_text 最前面一小段做摘要，真实测过这一小段只
# 覆盖了全文的 0.0025%、且全部来自最早的第一条记录，模型据此打出来的标签完全
# 不能代表这份文件实际装了什么。这两个问题本质是同一个病灶（记录被硬拼成一篇），
# 在提取这一步用均匀抽样解决，切分和打标签不需要知道发生过抽样，天然会拿到
# 跨度更广、更有代表性的内容。阈值不针对某个具体导出工具的 schema，任何"顶层
# 数组元素很多"的 JSON 都适用。
# ===== 代表性抽样工具包 =====
# 提取/切分/落库这几个阶段都会遇到"内容太多，需要瘦身但不能只看开头"这个
# 同一类问题，抽样算法只在这里写一份，谁需要谁 import，不各自重复实现。
#
# 早期版本里，_extract_json() 会对 JSON 顶层数组单独做"按记录抽样"（比如
# 1151 条聊天记录抽 300 条），设想是"记录级抽样比字符级抽样更公平"。真实
# 实测（用真实的 118MB conversations.json 对比）推翻了这个设想：按记录先
# 抽 300 条、再切分、再靠下面 chunk 数量上限抽 500 个 chunk，最终只覆盖了
# 1151 条原始记录里的 180 条(15.6%)；不做这道 JSON 专属预处理，直接对完整
# 原文切分、只靠 chunk 数量上限兜底抽样，反而覆盖了 362 条(31.5%)——因为
# "先按记录抽样"这一步会直接砍掉大多数记录（74%），砍掉的部分之后再也没有
# 机会被选中；而直接对完整原文做 chunk 级别的均匀抽样，抽样间隔天然比单条
# 记录的跨度更大，几乎每条原始记录都有平等机会被抽中。所以现在不再对 JSON
# 数组做任何格式专属的预处理，所有格式统一走"完整原文切分 → chunk 数量
# 上限兜底 → 落库前文本体积上限兜底"这一条路，没有例外分支。


def _sample_evenly(items, k):
    """从 items 里均匀抽样出最多 k 个，不是简单取前 k 个——保留跨越整个序列的
    代表性。只取前面的话，抽样结果只能反映序列最早的那一部分，看不出后面还有
    什么。
    """
    if len(items) <= k or k <= 0:
        return list(items)
    step = len(items) / k
    return [items[int(i * step)] for i in range(k)]


def spread_sample_text(text, max_len, pieces=5):
    """从很长的文本里均匀抽几段小样本拼起来，跨度覆盖全文，不是只读开头。

    跟 _sample_evenly() 解决的是同一类问题的另一种形态：_sample_evenly() 是
    "一份列表里挑几个元素"，这个函数是"一整段字符串里挑几个片段"，用于文本
    本身没有天然的"记录"分界（比如一整篇连贯文章），没法先转成列表再抽样的
    场景。跟单纯读开头相比，均匀分布在全文的样本才能真正反映内容的多样性，
    不会因为只读最前面一小段就得出片面的印象。
    """
    piece_len = max(max_len // pieces, 1)
    denom = max(pieces - 1, 1)
    samples = []
    for i in range(pieces):
        if len(text) <= piece_len:
            offset = 0
        else:
            offset = int(i * (len(text) - piece_len) / denom)
            offset = max(0, min(offset, len(text) - piece_len))
        samples.append(text[offset:offset + piece_len])
    return "\n...\n".join(samples).strip()


# 落库前对 extraction.text 做体积压缩时（见 indexing.py），会在压缩后的文本最
# 前面拼一句人话提示，如实说明"这不是完整原文"。格式定义只放这一处，写(indexing.py
# 生成这句提示)、读(tagging.py 打标签摘要要把它当噪声剥掉，不然它会被当成正文
# 内容读进去)两边共用，不各写一份——跟 SHEET_MARKER_PREFIX 是同一个思路。
_SAMPLE_NOTE_PATTERN = re.compile(r"^（提示：原文共 \d+ 字符，数量过多，[^）]*）\n?")


def format_sample_note(original_char_count):
    return f"（提示：原文共 {original_char_count} 字符，数量过多，这里是均匀抽样保留的代表性内容，不是完整原文）"


def strip_sample_note(text):
    """如果开头是 format_sample_note() 拼上去的那句提示，去掉它——这句提示是
    写给人看的元信息，不是文档正文，不该被当成内容片段读进摘要/embedding。
    没有这句提示的文本原样返回。
    """
    return _SAMPLE_NOTE_PATTERN.sub("", text, count=1)


def _extract_json(path):
    """把 JSON 重新格式化成带缩进的可读文本。选择保留完整结构（key 和 value 都在，
    不是只抽字符串值），因为 key 本身往往带有语义（比如 "name": "React" 里
    "name" 这个 key 是有意义的上下文，丢掉就退化成一堆孤立的字符串）。

    不在这里对内容量做任何抽样/截断——原因见上面"代表性抽样工具包"那段注释，
    多大的 JSON 都原样格式化返回，交给下游统一的 chunk 数量上限 + 落库体积
    上限兜底。
    """

    import json as json_module

    text = _decode_bytes(path.read_bytes())
    data = json_module.loads(text)
    return json_module.dumps(data, ensure_ascii=False, indent=2)


def _extract_pdf(path):
    """提取 PDF 的文字层；扫描版 PDF 没有文字层时返回空字符串。"""

    from pypdf import PdfReader

    reader = PdfReader(str(path))
    pages = []
    for page in reader.pages:
        pages.append(page.extract_text() or "")
    return "\n\n".join(pages)


# 从 WordDocument 内部数据流里，筛出连续 4 个字符以上、看起来像正常文字的片段
# （中文、英文字母、数字、常见标点、空白）。老版 .doc 文件里正文文字和大量格式/
# 结构相关的二进制控制信息交织在一起，这个正则本质是在"噪音里挑出可读的部分"。
_DOC_PRINTABLE_RUN_PATTERN = re.compile(r"[一-鿿A-Za-z0-9，。：；、（）,.:;()\s]{4,}")


def _extract_doc(path):
    """提取老版二进制 DOC（不是 DOCX，两者是不同格式，python-docx 只认新版 XML
    格式，读不了这种老格式）的正文文字。

    这不是完整、精确的 .doc 格式解析（真要做到那样，需要理解 Word 二进制格式里
    的 FIB/piece table 这些内部结构，工作量很大，跟当前"文本提取给检索/打标签用"
    的目标不成比例）。用的是一个真实测过有效的取巧办法：.doc 内部有一个专门存正文
    的 OLE 数据流叫 "WordDocument"，直接按 UTF-16LE 解码整个流，再用正则挑出
    连续的可打印字符片段——虽然会跳过表格/图注等复杂结构里的文字，段落顺序也不
    保证跟原文件完全一致，但对以纯文字为主的文档（这类文档占多数），实测能拿到
    干净、可读、内容真实的正文，足够支撑检索和打标签这两个用途。
    """

    import olefile

    ole = olefile.OleFileIO(str(path))
    try:
        raw = ole.openstream("WordDocument").read()
    finally:
        ole.close()

    text = raw.decode("utf-16-le", errors="ignore")
    runs = _DOC_PRINTABLE_RUN_PATTERN.findall(text)
    # \r（段落标记）/\x0b（手动换行）/\x0c（分页符）是 Word 内部的换行语义，
    # 统一换成真正的换行符，不然会跟其他空白字符混在一起，读不出原本的分段。
    cleaned = [re.sub(r"[\r\x0b\x0c]+", "\n", run).strip() for run in runs]
    return "\n".join(c for c in cleaned if c)


def _extract_docx(path):
    """提取 DOCX 段落和表格中的文字。"""

    from docx import Document

    document = Document(str(path))
    blocks = []

    for paragraph in document.paragraphs:
        if paragraph.text.strip():
            blocks.append(paragraph.text)

    # 表格不属于连续正文，但保留单元格之间的制表符，后续仍能检索到
    # 表格中的关键字；更复杂的表格结构化处理留到后续阶段。
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            if any(cells):
                blocks.append("\t".join(cells))

    return "\n".join(blocks)


def _extract_html(path):
    """去掉 HTML 中的脚本和样式，只保留可见文本。"""

    from bs4 import BeautifulSoup

    raw = path.read_bytes()
    soup = BeautifulSoup(raw, "html.parser")
    for element in soup(["script", "style", "noscript"]):
        element.decompose()
    return soup.get_text("\n")


# markdown 图片嵌入语法 ![描述](链接)——飞书/语雀这类笔记工具粘贴截图时常见，
# 链接通常是几十上百个字符的图床 CDN 地址（一长串没有任何空格的哈希/参数），
# 对语义检索/打标签没有价值，还会把切分算法逼到没有好断点、只能硬切的境地
# （真实抓到过 chunk 切在这种链接中间的案例）。整段清掉，不只是清掉链接部分——
# alt 文字通常也是"image.png"这种没有信息量的占位符，不是真实描述。
_MARKDOWN_IMAGE_PATTERN = re.compile(r"!\[[^\]]*\]\([^)]*\)")

# 普通 markdown 链接 [文字](链接)，跟图片嵌入不同：方括号里的文字通常是真实
# 有意义的内容（比如"详见 [React 官方文档](https://...)"里的"React 官方文档"），
# 值得保留，只清掉链接部分——所以只替换成方括号里的文字，不是整段删掉。
_MARKDOWN_LINK_PATTERN = re.compile(r"\[([^\]]*)\]\([^)]*\)")


def _normalize_text(text):
    """做最小清洗，保留正文顺序，不做语义改写。"""

    text = _MARKDOWN_IMAGE_PATTERN.sub("", text)
    text = _MARKDOWN_LINK_PATTERN.sub(r"\1", text)

    lines = [line.strip() for line in text.replace("\r\n", "\n").replace("\r", "\n").split("\n")]

    # 连续空行压缩成一行，避免 HTML/PDF 的排版空白污染后续分段。
    normalized_lines = []
    previous_blank = False
    for line in lines:
        is_blank = not line
        if is_blank and previous_blank:
            continue
        normalized_lines.append(line)
        previous_blank = is_blank

    return "\n".join(normalized_lines).strip()
